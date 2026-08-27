import os
import re
import fitz  # PyMuPDF
import base64
import json

try:
    import docx
except ImportError:
    docx = None
try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
except ImportError:
    try:
        from langchain.text_splitter import RecursiveCharacterTextSplitter
    except ImportError:
        RecursiveCharacterTextSplitter = None

from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain_community.vectorstores import FAISS
from langchain_core.prompts import PromptTemplate
from langchain_core.documents import Document

MAX_PAGES_PER_DOC = 150

def get_api_key():
    """Retrieve Google Gemini API Key from environment."""
    key = os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY", "")
    return key.strip('"').strip("'").strip()

def sanitize_html_output(raw_html: str) -> str:
    """Sanitizes AI-generated HTML to prevent XSS payloads while preserving clean formatting."""
    if not raw_html:
        return ""
    # Strip script tags, iframes, object/embeds, and inline event handlers
    cleaned = re.sub(r'<\s*script[^>]*>.*?<\s*/\s*script\s*>', '', raw_html, flags=re.DOTALL | re.IGNORECASE)
    cleaned = re.sub(r'<\s*iframe[^>]*>.*?<\s*/\s*iframe\s*>', '', cleaned, flags=re.DOTALL | re.IGNORECASE)
    cleaned = re.sub(r'<\s*object[^>]*>.*?<\s*/\s*object\s*>', '', cleaned, flags=re.DOTALL | re.IGNORECASE)
    cleaned = re.sub(r'<\s*embed[^>]*>.*?', '', cleaned, flags=re.DOTALL | re.IGNORECASE)
    cleaned = re.sub(r'on\w+\s*=\s*["\'][^"\']*["\']', '', cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r'javascript\s*:', '', cleaned, flags=re.IGNORECASE)
    return cleaned

def process_documents(file_paths, max_pages: int = 150):
    """
    Extracts text and embedded links from PDFs, Word docs, and text files.
    Splits text into chunks for vector indexing with DoS/Bomb and Subscription Tier page protections.
    """
    documents = []
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=4000, chunk_overlap=400)
    
    for file_path in file_paths:
        try:
            ext = os.path.splitext(file_path)[1].lower()
            
            if ext == '.pdf':
                doc = fitz.open(file_path)
                doc_len = len(doc)
                if doc_len > max_pages:
                    doc.close()
                    raise ValueError(
                        f"PAGE_LIMIT_EXCEEDED: File '{os.path.basename(file_path)}' contains {doc_len} pages, "
                        f"which exceeds your plan limit of {max_pages} pages. Upgrade to ZenDoc Pro to unlock up to 500 pages."
                    )
                
                total_pages = doc_len
                for page_num in range(total_pages):
                    page = doc.load_page(page_num)
                    
                    # Extract text blocks preserving tabular structure & layout
                    blocks = page.get_text("blocks")
                    page_text_blocks = [b[4].strip() for b in blocks if len(b) > 4 and b[4].strip()]
                    text = "\n\n".join(page_text_blocks) if page_text_blocks else page.get_text("text")
                    
                    # Extract embedded hyperlinks so Gemini can ground links
                    links = page.get_links()
                    extracted_links = [
                        link.get("uri") for link in links 
                        if link.get("kind") == fitz.LINK_URI and link.get("uri")
                    ]
                    if extracted_links:
                        text += "\n\n[Hyperlinks in Page " + str(page_num + 1) + "]: " + ", ".join(extracted_links)
                        
                    if text and text.strip():
                        chunks = text_splitter.split_text(text)
                        for chunk in chunks:
                            documents.append(
                                Document(
                                    page_content=f"[Document: {os.path.basename(file_path)} | Page {page_num + 1}]\n{chunk}", 
                                    metadata={"source": file_path, "filename": os.path.basename(file_path), "page": page_num}
                                )
                            )
                doc.close()
                
            elif ext == '.docx':
                if not docx:
                    raise ValueError("DOCX parsing library is not installed on this host. Please upload a PDF or TXT file, or install python-docx.")
                doc = docx.Document(file_path)
                text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                if text.strip():
                    chunks = text_splitter.split_text(text)
                    for chunk in chunks:
                        documents.append(
                            Document(
                                page_content=chunk,
                                metadata={"source": file_path, "filename": os.path.basename(file_path), "page": -1}
                            )
                        )
                        
            elif ext == '.txt':
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
                if text.strip():
                    chunks = text_splitter.split_text(text)
                    for chunk in chunks:
                        documents.append(
                            Document(
                                page_content=chunk,
                                metadata={"source": file_path, "filename": os.path.basename(file_path), "page": -1}
                            )
                        )
                        
        except Exception as e:
            if "PAGE_LIMIT_EXCEEDED" in str(e):
                raise
            print(f"Error reading {file_path}: {e}")
            
    return documents

def get_embeddings_model():
    """Initializes Google GenAI Embeddings with automatic fallback."""
    api_key = get_api_key()
    if not api_key:
        raise ValueError("Google Gemini API Key is missing. Please configure GEMINI_API_KEY in environment variables.")
    try:
        return GoogleGenerativeAIEmbeddings(model="models/gemini-embedding-2", google_api_key=api_key)
    except Exception:
        return GoogleGenerativeAIEmbeddings(model="models/gemini-embedding-001", google_api_key=api_key)

def get_vector_store(documents):
    """Embeds documents into an in-memory & persisted FAISS vector index."""
    embeddings = get_embeddings_model()
    vector_store = FAISS.from_documents(documents, embedding=embeddings)
    vector_store.save_local("./faiss_index")
    return vector_store

def get_conversational_chain():
    """Builds the Gemini 2.5 Flash multimodal conversational reasoning chain."""
    prompt_template = """
    You are ZenDoc AI, an enterprise document intelligence and reasoning engine developed by Aneevarp Solutions, powered by Google Gemini 2.5 Flash.

    Answer the user's question accurately, concisely, and insightfully based ONLY on the provided document context.

    Follow these strict formatting standards:
    - Use <b>bold text</b> for key metrics, conclusions, terms, and critical data points.
    - Use clean bullet points (<ul><li>...</li></ul>) or numbered lists (<ol><li>...</li></ol>) to break down multi-step answers.
    - Render tabular data or financial balance sheets in clean HTML tables (<table border="1">...</table>) or clean markdown.
    - If comparing multiple documents, explicitly state the source filename for each point (e.g. <b>[Document A.pdf]</b> vs <b>[Document B.pdf]</b>).
    - If links/URLs are found in the context, render them as clickable links: <a href="..." target="_blank" rel="noopener">link</a>.
    - If the context does NOT contain the answer, politely respond: "I could not find relevant information in the uploaded document. Please check the file or rephrase your question." Do NOT hallucinate.

    CRITICAL: Output your response as a valid JSON object ONLY, with no surrounding markdown code blocks (no ```json).
    JSON Structure:
    {{
      "answer": "<HTML-formatted comprehensive answer>",
      "exact_quote": "<Exact short verbatim quote from the text supporting this answer, or empty string if not found>"
    }}

    Document Context:
    {context}

    User Question:
    {question}

    JSON Output:
    """
    
    api_key = get_api_key()
    if not api_key:
        raise ValueError("Google Gemini API Key is missing. Please set GEMINI_API_KEY.")

    model = ChatGoogleGenerativeAI(
        model="gemini-2.5-flash", 
        temperature=0.2, 
        google_api_key=api_key,
        model_kwargs={"response_mime_type": "application/json"}
    )
    prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
    return prompt | model

def get_base64_image(pdf_path, page_num, quote=""):
    """Crops visual bounding-box snippet from original PDF page for visual grounding."""
    try:
        if not os.path.exists(pdf_path):
            return None
            
        doc = fitz.open(pdf_path)
        if page_num < 0 or page_num >= len(doc):
            doc.close()
            return None
            
        page = doc.load_page(page_num)
        clip_rect = None
        
        if quote and len(quote.strip()) > 3:
            words = quote.strip().split()
            search_phrase = " ".join(words[:8]) if len(words) > 8 else quote.strip()
            rects = page.search_for(search_phrase)
            
            if not rects and len(words) > 5:
                search_phrase = " ".join(words[-8:])
                rects = page.search_for(search_phrase)
                
            if rects:
                x0 = min([r.x0 for r in rects])
                y0 = min([r.y0 for r in rects])
                x1 = max([r.x1 for r in rects])
                y1 = max([r.y1 for r in rects])
                
                # Expand bounding box to capture surrounding visual context
                padding_y = 50
                padding_x = 30
                page_rect = page.rect
                
                clip_rect = fitz.Rect(
                    max(0, x0 - padding_x),
                    max(0, y0 - padding_y),
                    min(page_rect.x1, x1 + padding_x),
                    min(page_rect.y1, y1 + padding_y)
                )

        pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), clip=clip_rect)
        img_bytes = pix.tobytes("png")
        b64_img = base64.b64encode(img_bytes).decode("utf-8")
        doc.close()
        return b64_img
    except Exception as e:
        print(f"Error rendering visual bounding box: {e}")
        return None

def process_user_query(user_question):
    """Performs semantic vector search in FAISS and invokes Gemini conversational reasoning."""
    if not os.path.exists("./faiss_index"):
        return {
            "answer": "<b>No active document found.</b> Please upload a PDF, Word, or Text document first.",
            "source_image": None,
            "page": None,
            "file_type": None,
            "filename": None
        }
        
    embeddings = get_embeddings_model()
    db = FAISS.load_local("./faiss_index", embeddings, allow_dangerous_deserialization=True)
    
    # Dynamic Top-K chunk retrieval: Expand for broad / analytical questions
    q_lower = user_question.lower()
    is_broad_query = any(w in q_lower for w in ["summary", "summarize", "overview", "all", "table", "financial", "compare", "difference", "metrics", "risks", "key terms", "action items", "dossier"])
    k_val = 8 if is_broad_query else 4
    
    docs = db.similarity_search(user_question, k=k_val)
    if not docs:
        return {
            "answer": "No relevant context found in the uploaded document for your query.",
            "source_image": None,
            "page": None,
            "file_type": None,
            "filename": None
        }
        
    context_text = "\n\n---\n\n".join([doc.page_content for doc in docs])
    chain = get_conversational_chain()
    
    response = chain.invoke({"context": context_text, "question": user_question})
    
    # Parse JSON output safely
    response_text = response.content.strip()
    if response_text.startswith("```json"):
        response_text = response_text[7:]
    if response_text.endswith("```"):
        response_text = response_text[:-3]
    response_text = response_text.strip()
    
    try:
        data = json.loads(response_text)
        answer = data.get("answer", "Could not format answer.")
        exact_quote = data.get("exact_quote", "")
    except Exception as e:
        print("JSON parse error:", e)
        answer = response_text
        exact_quote = ""
    
    # Visual Grounding extraction: Search through retrieved docs for best visual crop
    source_image = None
    best_doc = docs[0]
    source_file = best_doc.metadata.get("source", "")
    filename = best_doc.metadata.get("filename", os.path.basename(source_file) if source_file else "")
    page_num = best_doc.metadata.get("page")
    file_ext = os.path.splitext(source_file)[1].lower() if source_file else ""
    
    # If the first doc is a PDF, crop it; otherwise check if any retrieved doc is a PDF with visual proof
    for d in docs:
        s_file = d.metadata.get("source", "")
        p_num = d.metadata.get("page")
        if s_file and os.path.splitext(s_file)[1].lower() == '.pdf' and p_num is not None and p_num >= 0:
            cropped = get_base64_image(s_file, p_num, exact_quote)
            if cropped:
                source_image = cropped
                source_file = s_file
                filename = d.metadata.get("filename", os.path.basename(s_file))
                page_num = p_num
                file_ext = '.pdf'
                break
    
    return {
        "answer": sanitize_html_output(answer),
        "source_image": source_image,
        "page": page_num,
        "file_type": file_ext,
        "filename": filename,
        "exact_quote": exact_quote
    }
